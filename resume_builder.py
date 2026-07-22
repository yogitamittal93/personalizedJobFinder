import os
import re
import json
import time
import requests
from bs4 import BeautifulSoup
from dotenv import load_dotenv
from gemini_client import llm

load_dotenv()

TAILORED_RESUMES_DIR = os.path.join(os.path.dirname(__file__), "tailored_resumes")

# ─────────────────────────────────────────────
# CANDIDATE STATIC DATA
# Only name, contact, links are hardcoded here.
# Everything else is dynamically tailored per job.
# ─────────────────────────────────────────────
CANDIDATE_HEADER = {
    "name": "Yogita Singla",
    "email": "yogitamittal.tech@gmail.com",
    "portfolio": "https://portfolio-three-sigma-mp0vvhcq3h.vercel.app/",
    "github": "https://github.com/yogitamittal93",
    "linkedin": "https://linkedin.com/in/yogitamittal",
    "location": "Remote / India"
}

# ─────────────────────────────────────────────
# FULL EXPERIENCE BANK
# Every real experience, project, and skill you have.
# Gemini picks, suppresses, and reframes from this bank
# based on what each specific job actually needs.
# NEVER fabricate — only reframe what exists here.
# ─────────────────────────────────────────────
EXPERIENCE_BANK = """
=== WORK EXPERIENCE ===

[Measured Inc. — Senior Full Stack Developer / Marketing Tech Lead]
Duration: ~2 years
Stack: HubSpot CRM, HubSpot Workflows, API integrations, JavaScript, Node.js, PostgreSQL
Key work:
- Architected and implemented full HubSpot CRM integration including custom properties,
  workflow automation, and lifecycle stage triggers
- Built behavior-triggered re-engagement email flows covering 500k+ contacts
- Integrated GTM (Google Tag Manager) for conversion tracking across multi-brand properties
- Led performance optimization initiatives reducing page load by 40%
- Worked directly with marketing team to translate campaign goals into technical automation

[Gardens Alive / Spring Hill Nursery / Gurney's — Lead Developer & Email Marketing Tech]
Duration: 5+ years
Stack: PHP, JavaScript, Commercev3, Klaviyo, email automation, A/B testing, segmentation
Key work:
- Led full ecommerce platform ownership on Commercev3 for multi-brand horticulture portfolio
  (Spring Hill Nursery, Gurney's, and 4 sister brands)
- Built behavior-triggered Klaviyo email automation: browse abandonment, cart recovery,
  win-back flows, post-purchase sequences — contributed to shifting client budget from
  print catalogs to digital (signature career achievement)
- Developed custom PHP integrations between ecommerce platform and email service providers
- Implemented advanced segmentation: RFM scoring, purchase history, engagement tiers
- Managed A/B testing programs across email subject lines, send times, CTAs
- Built custom reporting dashboards for campaign performance

[Lakshmi Iron Company — Technical Lead]
Stack: WordPress, WooCommerce, PHP, JavaScript, custom plugins
Key work:
- Built full WooCommerce store with custom product configurator
- Developed custom WordPress plugins for inventory management and order processing
- Implemented performance optimization (caching, CDN, image optimization)

[Freelance — Full Stack Developer (ongoing)]
Stack: Next.js, Node.js, NestJS, TypeScript, PostgreSQL, Prisma, React, Tailwind CSS,
       Capacitor (iOS/Android), Stripe, OAuth (Google/GitHub via Passport.js)
Key work:
- Built and deployed MomDigital: maternal health app with NestJS backend, Next.js frontend,
  Python Flask ML service, RAG pipeline (ChromaDB + Groq Llama 3.3 + sentence-transformers),
  PostgreSQL via Prisma — full solo architecture and build
- Resolved 157 TypeScript compilation errors across large codebase
- Wrapped client web app (Lovable/TanStack) into native Android + iOS apps using Capacitor,
  tested on real Android device
- Integrated Google OAuth and GitHub OAuth via Passport.js
- Built community social feed, exercise log with SVG calendar heatmap and progress rings,
  trusted-helper nanny/chef checklists with scoring, baby food timeline
- Implemented Stripe payment integrations for client projects
- Set up structured payment protection using staged delivery and JKS keystore

=== SKILLS ===

Languages: PHP (13 years), JavaScript (13 years), TypeScript, Python, SQL
Frameworks: Next.js, Node.js, NestJS, React, Flask, Express, WordPress, WooCommerce
Databases: PostgreSQL, MySQL, SQLite, ChromaDB (vector DB)
ORM / Query: Prisma, raw SQL
Email & MarTech: Klaviyo, HubSpot, Attentive, GTM, Mailchimp, ActiveCampaign
  — lifecycle flows, segmentation, A/B testing, deliverability, behavioral triggers
eCommerce: Commercev3, WooCommerce, Shopify (familiar), headless commerce patterns,
  product pages, cart, checkout, order management, inventory
APIs & Integrations: REST, webhooks, OAuth 2.0, Stripe, Passport.js, third-party SaaS integrations
AI / ML: RAG pipeline architecture, ChromaDB, Groq (Llama 3.3), sentence-transformers,
  Python Flask ML service, prompt engineering, vector embeddings
Mobile: Capacitor (iOS + Android), PWA
DevOps / Infra: GitHub Actions, Docker (familiar), Vercel, Netlify, basic Linux CLI
CMS: WordPress (custom themes + plugins), headless CMS patterns, Contentful (familiar)
Analytics & Tracking: Google Tag Manager, Google Analytics, conversion tracking, funnel analysis
Performance: Core Web Vitals optimization, caching strategies, CDN, image optimization
Tools: Git, GitHub, Figma (read), Jira, Linear, Notion

=== PROJECTS ===

[MomDigital — Maternal & Infant Health App]
Tech: NestJS, Next.js, Python Flask, PostgreSQL, Prisma, ChromaDB, Groq, sentence-transformers,
      Google OAuth, GitHub OAuth, Tailwind CSS, Docker
Description: Full-stack maternal health platform targeting Indian mothers. Integrates MBBS
clinical guidelines with Ayurvedic wisdom via a RAG pipeline. Features: community social feed,
exercise tracking with SVG visualizations, appointment booking, trusted-helper checklists,
baby food timeline, AI chatbot powered by Llama 3.3 via Groq. Vector database (634MB)
hosted on Hugging Face Hub. Planned: multi-tenant expert access, freemium model,
WhatsApp/Jitsi live sessions.

[Capacitor Mobile App — Client Project]
Tech: Capacitor, Android, iOS, TanStack Start, Google OAuth
Description: Wrapped client's Lovable/TanStack web app into native Android and iOS apps.
Successfully tested on real Android device. Structured staged payment and JKS keystore
delivery for client protection.

[AI Job Application Agent — PersonalizedJobFinder]
Tech: Python, SQLite, Playwright, Gemini API, Greenhouse/Lever/Ashby APIs, SMTP
Description: Automated job discovery pipeline that scrapes 5 job board sources,
evaluates listings against candidate profile using Gemini LLM, deduplicates,
and sends HTML digest emails with match scoring and interview prep plans.

=== EDUCATION ===
[Add your actual degree, institution, and year here]

=== SOFT SKILLS & WORKING STYLE ===
- 13 years of client delivery across industries — reliable, deadline-driven
- Strong communicator between technical and non-technical stakeholders
- Self-directed: built entire MomDigital stack solo while managing freelance clients
- Comfortable wearing multiple hats: developer, architect, marketer, analyst
- Remote-first, async-first work style
- Deeply curious — picks up new stacks and frameworks quickly (NestJS, Capacitor,
  RAG pipelines all self-taught in the last 12 months)
"""


# ─────────────────────────────────────────────
# WORD DOCUMENT BUILDER
# Converts structured plain-text resume sections
# from Gemini into a properly formatted .docx
# ─────────────────────────────────────────────
def build_docx(sections: dict, company: str, title: str, docx_path: str) -> str:
    """
    Takes parsed resume sections dict and writes an ATS-safe Word document.
    sections keys: summary, skills, experience, projects, education
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── Page margins (narrow) ────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin   = Pt(46)
        section.right_margin  = Pt(46)

    # ── Default style ────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    def add_section_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)  # brand blue
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(2)
        # Add a bottom border via paragraph border (simple rule)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '1d4ed8')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_bullet(text):
        # Remove leading hyphens/bullets from Gemini output
        clean = text.lstrip('-•* ').strip()
        if not clean:
            return
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent  = Pt(14)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(clean)
        run.font.size = Pt(10)

    # ── Name header ─────────────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(CANDIDATE_HEADER['name'])
    name_run.bold = True
    name_run.font.size = Pt(18)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_before = Pt(2)
    contact_p.paragraph_format.space_after  = Pt(6)
    contact_text = (
        f"{CANDIDATE_HEADER['email']}  |  "
        f"{CANDIDATE_HEADER['portfolio']}  |  "
        f"{CANDIDATE_HEADER['github']}  |  "
        f"{CANDIDATE_HEADER['location']}"
    )
    contact_run = contact_p.add_run(contact_text)
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    # ── Professional Summary ─────────────────────────────────────────────────
    if sections.get('summary'):
        add_section_heading('Professional Summary')
        p = doc.add_paragraph(sections['summary'].strip())
        p.paragraph_format.space_after = Pt(4)

    # ── Core Skills ─────────────────────────────────────────────────────────
    if sections.get('skills'):
        add_section_heading('Core Skills')
        for line in sections['skills'].strip().splitlines():
            line = line.strip()
            if not line:
                continue
            if ':' in line:
                # "Languages: PHP, JS" → bold label, normal text
                label, _, rest = line.partition(':')
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after  = Pt(1)
                lbl_run = p.add_run(label.strip() + ': ')
                lbl_run.bold = True
                lbl_run.font.size = Pt(10)
                p.add_run(rest.strip()).font.size = Pt(10)
            else:
                add_bullet(line)

    # ── Professional Experience ──────────────────────────────────────────────
    if sections.get('experience'):
        add_section_heading('Professional Experience')
        for line in sections['experience'].strip().splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith('-') or stripped.startswith('•'):
                add_bullet(stripped)
            else:
                # Company/title line — bold it
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(1)
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(10)

    # ── Key Projects ─────────────────────────────────────────────────────────
    if sections.get('projects'):
        add_section_heading('Key Projects')
        for line in sections['projects'].strip().splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith('-') or stripped.startswith('•'):
                add_bullet(stripped)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after  = Pt(1)
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(10)

    # ── Education ────────────────────────────────────────────────────────────
    if sections.get('education'):
        add_section_heading('Education')
        for line in sections['education'].strip().splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith('-') or stripped.startswith('•'):
                add_bullet(stripped)
            else:
                p = doc.add_paragraph(stripped)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(1)

    # ── Footer ───────────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(12)
    footer_run = footer_p.add_run(
        f"Tailored for {company} — {title}  |  Generated by JobCraft AI"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)

    doc.save(docx_path)
    print(f"✅ Word resume written: {docx_path}")
    return docx_path


# ─────────────────────────────────────────────
# COMPANY CULTURE SCRAPER
# Fetches About/Careers page to extract culture
# signals, mission language, and values — so the
# resume and summary mirror the company's own words
# ─────────────────────────────────────────────
def scrape_company_culture(company: str, job_url: str) -> str:
    """
    Attempts to fetch the company's About or Careers page to extract
    culture signals, mission statement, and values language.
    Falls back gracefully if scraping fails.
    """
    culture_text = ""
    headers = {"User-Agent": "Mozilla/5.0"}

    # Derive base domain from job URL
    try:
        from urllib.parse import urlparse
        parsed = urlparse(job_url)
        # Handle ATS-hosted URLs (jobs.lever.co/company, boards.greenhouse.io/company)
        if any(ats in parsed.netloc for ats in ["lever.co", "greenhouse.io", "ashbyhq.com"]):
            # Try to guess company website from the company name
            company_slug = re.sub(r'\W+', '', company).lower()
            candidate_urls = [
                f"https://www.{company_slug}.com/about",
                f"https://www.{company_slug}.com/careers",
                f"https://{company_slug}.com/about",
            ]
        else:
            base = f"{parsed.scheme}://{parsed.netloc}"
            candidate_urls = [
                f"{base}/about",
                f"{base}/careers",
                f"{base}/company",
            ]

        for url in candidate_urls:
            try:
                res = requests.get(url, timeout=6, headers=headers)
                if res.status_code == 200:
                    soup = BeautifulSoup(res.text, 'html.parser')
                    for el in soup(["script", "style", "nav", "footer", "header"]):
                        el.extract()
                    text = " ".join(soup.get_text().split())
                    if len(text) > 200:  # meaningful content found
                        culture_text = text[:3000]  # cap to keep prompt manageable
                        print(f"   🏢 Company culture scraped from: {url}")
                        break
            except Exception:
                continue

    except Exception as e:
        print(f"   ⚠️  Culture scraping failed: {e}")

    if not culture_text:
        print(f"   ℹ️  No culture page found for {company} — proceeding without it.")

    return culture_text


# ─────────────────────────────────────────────
# JD KEYWORD EXTRACTOR
# ─────────────────────────────────────────────
def extract_jd_keywords(job_description: str) -> list:
    """
    Extracts tech keywords from the JD for ATS mirroring.
    Returns up to 40 keywords.
    """
    tech_pattern = re.compile(
        r'\b(React|Next\.js|Node\.js|NestJS|TypeScript|JavaScript|Python|PHP|Laravel|'
        r'WordPress|WooCommerce|HubSpot|Klaviyo|Shopify|Commercev3|PostgreSQL|MySQL|'
        r'MongoDB|Redis|Docker|Kubernetes|AWS|GCP|Azure|GraphQL|REST(?:ful)?|API|'
        r'CI\/CD|Git(?:Hub)?|Vercel|Netlify|Supabase|TailwindCSS|Tailwind|'
        r'Vue\.js|Angular|Svelte|Express|Prisma|Stripe|Figma|Jest|Cypress|'
        r'A\/B\s?testing|SEO|CMS|E-?commerce|SaaS|B2B|B2C|'
        r'microservice|serverless|full.?stack|front.?end|back.?end|'
        r'DevOps|Agile|Scrum|GTM|Google Tag Manager|Klaviyo|Attentive|'
        r'lifecycle|segmentation|automation|webhook|OAuth|Capacitor|'
        r'ChromaDB|RAG|LLM|machine learning|Prisma|tRPC|Playwright)\b',
        re.IGNORECASE
    )
    keywords = list(set(tech_pattern.findall(job_description)))
    return keywords[:40]


# ─────────────────────────────────────────────
# ROLE TYPE DETECTOR
# Routes resume shape based on job type so the
# emphasis shifts appropriately per role category
# ─────────────────────────────────────────────
def detect_role_type(title: str, description: str) -> str:
    """
    Returns a role type string that guides resume emphasis.
    """
    combined = (title + " " + description).lower()

    if any(kw in combined for kw in [
        "growth engineer", "growth hacking", "lifecycle", "retention",
        "activation", "klaviyo", "braze", "iterable", "email engineer"
    ]):
        return "GROWTH_ENGINEER"

    if any(kw in combined for kw in [
        "founding engineer", "first engineer", "0 to 1", "seed stage",
        "generalist", "full-stack generalist", "technical co-founder",
        "wear many hats", "small team"
    ]):
        return "FOUNDING_ENGINEER"

    if any(kw in combined for kw in [
        "ecommerce", "e-commerce", "shopify", "woocommerce", "dtc",
        "direct to consumer", "checkout", "cart", "product catalog",
        "headless commerce", "commercetools"
    ]):
        return "ECOMMERCE_ENGINEER"

    if any(kw in combined for kw in [
        "wordpress", "wp ", "wp-", "theme", "plugin", "cms",
        "headless wordpress", "elementor", "gutenberg"
    ]):
        return "WORDPRESS_ENGINEER"

    if any(kw in combined for kw in [
        "integration", "hubspot", "crm", "api integration",
        "martech", "marketing technology", "marketing platform"
    ]):
        return "MARTECH_INTEGRATION"

    if any(kw in combined for kw in [
        "machine learning", "ml engineer", "ai engineer", "llm",
        "rag", "vector", "embeddings", "fine-tuning", "data science"
    ]):
        return "ML_ADJACENT"

    return "FULLSTACK_GENERAL"


# ─────────────────────────────────────────────
# ROLE-SPECIFIC EMPHASIS GUIDE
# Tells Gemini exactly what to lead with and what
# to push to the background per role type
# ─────────────────────────────────────────────
ROLE_EMPHASIS = {
    "GROWTH_ENGINEER": """
RESUME EMPHASIS FOR THIS ROLE TYPE (Growth Engineer):
LEAD WITH: Klaviyo automation at Gardens Alive — behavior-triggered flows, segmentation,
           A/B testing, lifecycle campaigns, shifting client from print to digital.
           HubSpot at Measured — workflow automation, lifecycle triggers, GTM tracking.
SECOND: Full stack skills (Next.js, Node.js) as the technical foundation that lets you
        build what the marketing team needs, not just configure tools.
SUPPRESS / MINIMIZE: MomDigital ML/RAG details, Capacitor mobile work, deep backend architecture.
SUMMARY TONE: "I sit at the intersection of engineering and growth — I've spent 13 years
building the automation systems that drive revenue, not just the products."
""",

    "FOUNDING_ENGINEER": """
RESUME EMPHASIS FOR THIS ROLE TYPE (Founding / Generalist Engineer):
LEAD WITH: Breadth — PHP+JS (13 years), email automation, ecommerce, Next.js, NestJS,
           Python, mobile (Capacitor), AI/RAG pipeline — all self-taught and shipped.
           MomDigital as proof: solo-built a full production stack with 5 different services.
SECOND: Client delivery track record — 13 years of shipping for real businesses.
SUPPRESS / MINIMIZE: Nothing — generalist roles want to see everything. Include it all.
SUMMARY TONE: "I'm the engineer who ships features in the morning, fixes infrastructure
at noon, and talks to users in the afternoon. That's not a gap — that's the job."
""",

    "ECOMMERCE_ENGINEER": """
RESUME EMPHASIS FOR THIS ROLE TYPE (eCommerce Engineer):
LEAD WITH: 5+ years on Commercev3 across 6 brands (Gardens Alive portfolio) —
           product pages, cart, checkout, order management, inventory, integrations.
           WooCommerce builds (Lakshmi Iron). Klaviyo email automation tied to purchase events.
SECOND: PHP depth (13 years), JavaScript, custom integrations, performance optimization.
SUPPRESS / MINIMIZE: MomDigital ML/RAG pipeline (not relevant), Capacitor mobile.
SUMMARY TONE: "eCommerce isn't just a stack to me — I've owned it end to end for
multi-brand portfolios, from product catalog to post-purchase lifecycle."
""",

    "WORDPRESS_ENGINEER": """
RESUME EMPHASIS FOR THIS ROLE TYPE (WordPress Engineer):
LEAD WITH: WordPress custom theme and plugin development (Lakshmi Iron + freelance).
           PHP depth (13 years), WooCommerce, performance optimization, custom post types,
           hooks and filters, REST API integration with WordPress.
SECOND: JavaScript, Next.js (headless WordPress patterns), GTM, page speed.
SUPPRESS / MINIMIZE: MomDigital ML/RAG, Capacitor mobile, NestJS/Python backend.
SUMMARY TONE: "WordPress isn't a limitation to me — it's a platform I know deeply
enough to build anything on, from simple sites to complex ecommerce ecosystems."
""",

    "MARTECH_INTEGRATION": """
RESUME EMPHASIS FOR THIS ROLE TYPE (MarTech / Integration Engineer):
LEAD WITH: HubSpot full integration at Measured (CRM, workflows, lifecycle triggers).
           Klaviyo at Gardens Alive (behavioral flows, segmentation, API integration).
           GTM implementation, webhook handling, third-party SaaS integrations.
SECOND: Full stack skills (Node.js, PHP, PostgreSQL) as the engine behind integrations.
SUPPRESS / MINIMIZE: MomDigital ML/AI work, mobile Capacitor, deep database architecture.
SUMMARY TONE: "I've connected the tools that drive revenue — HubSpot, Klaviyo, GTM —
and I've built the custom glue code that makes them actually work together."
""",

    "ML_ADJACENT": """
RESUME EMPHASIS FOR THIS ROLE TYPE (ML-adjacent / AI Engineer):
LEAD WITH: MomDigital RAG pipeline — ChromaDB vector DB (634MB), Groq Llama 3.3,
           sentence-transformers, Python Flask ML service, prompt engineering.
           Python fluency, PostgreSQL, REST API design.
SECOND: Full stack foundation (Next.js, NestJS, Node.js) — you can ship the product,
        not just the model. Practical AI implementation, not research.
SUPPRESS / MINIMIZE: PHP/WooCommerce/email marketing work — not relevant here.
SUMMARY TONE: "I don't just integrate AI APIs — I built the RAG pipeline, the vector
store, and the product around it. Practical AI, production-ready."
""",

    "FULLSTACK_GENERAL": """
RESUME EMPHASIS FOR THIS ROLE TYPE (Full Stack General):
LEAD WITH: 13 years of shipping across diverse stacks — PHP, JavaScript, TypeScript,
           Next.js, Node.js, NestJS, Python, PostgreSQL.
           MomDigital as the flagship project: solo-architected and built.
SECOND: Commercial track record — ecommerce, email automation, client delivery.
SUPPRESS / MINIMIZE: Nothing critical to hide — show breadth but keep bullets tight.
SUMMARY TONE: "13 years of full-stack delivery across ecommerce, martech, and product
engineering. I pick up new stacks fast and ship things that work."
"""
}


# ─────────────────────────────────────────────
# MAIN RESUME GENERATOR
# Outputs a .docx Word file + .md alignment audit
# ─────────────────────────────────────────────
def generate_tailored_resume(job: dict, profile: dict) -> tuple:
    """
    Generates a tailored, ATS-optimized Word (.docx) resume AND alignment audit markdown.

    The resume is shaped per role type — a growth engineering role gets a completely
    different emphasis ordering than a WordPress role or a founding engineer role.
    Company culture language is scraped and mirrored in the summary.
    Every bullet comes from the real experience bank. Nothing fabricated.
    Returns (docx_path, md_path).
    """
    if not os.path.exists(TAILORED_RESUMES_DIR):
        os.makedirs(TAILORED_RESUMES_DIR)

    company = job.get('company', 'Company')
    title   = job.get('title', 'Role')
    company_clean = re.sub(r'\W+', '', company).capitalize()
    title_clean   = re.sub(r'\W+', '', title).capitalize()

    docx_path = os.path.join(TAILORED_RESUMES_DIR, f"{company_clean}_{title_clean}_Resume.docx")
    md_path   = os.path.join(TAILORED_RESUMES_DIR, f"{company_clean}_{title_clean}_Alignment.md")

    # Always regenerate fresh — never serve stale
    for path in [docx_path, md_path]:
        if os.path.exists(path):
            try:
                os.remove(path)
            except Exception:
                pass

    job_description = job.get('description', '')
    profile_str     = json.dumps(profile, indent=2)

    # ── Enrich context ───────────────────────────────────────────────────────
    jd_keywords   = extract_jd_keywords(job_description)
    role_type     = detect_role_type(title, job_description)
    role_emphasis = ROLE_EMPHASIS.get(role_type, ROLE_EMPHASIS["FULLSTACK_GENERAL"])
    culture_text  = scrape_company_culture(company, job.get('url', ''))

    keywords_str  = ", ".join(jd_keywords) if jd_keywords else "see JD above"
    culture_block = f"\n=== COMPANY CULTURE & MISSION (mirror this language) ===\n{culture_text}" \
                    if culture_text else "\n(No culture page found — infer tone from JD language.)"

    print(f"   🎯 Role type detected: {role_type}")
    print(f"   📌 {len(jd_keywords)} JD keywords to mirror")

    # ── Master prompt ────────────────────────────────────────────────────────
    prompt = f"""
You are the Elite Resume Tailoring Engine of JobCraft AI.
Your SOLE objective: make Yogita Singla the most compelling candidate for this specific
role — both to ATS systems and to the human hiring manager who reads it.

=== TARGET POSITION ===
Company: {company}
Role: {title}
Job Description:
{job_description}

{culture_block}

=== CANDIDATE EXPERIENCE BANK (ONLY source of truth — never fabricate) ===
{EXPERIENCE_BANK}

=== CANDIDATE PROFILE JSON (supplementary context) ===
{profile_str}

=== ROLE TYPE DETECTED: {role_type} ===
{role_emphasis}

=== JD KEYWORDS TO MIRROR FOR ATS (use these exact strings in resume) ===
{keywords_str}

=== CRITICAL RULES — VIOLATIONS WILL INVALIDATE THE OUTPUT ===

RULE 1 — NO FABRICATION:
Every skill, metric, company name, project, and technology in the resume MUST exist
verbatim in the Experience Bank above. Do NOT invent numbers, tools, or responsibilities.
If a JD requirement genuinely doesn't exist in the bank, acknowledge the gap in the
audit — do NOT insert it into the resume.

RULE 2 — ATS KEYWORD MIRRORING:
Every keyword from the JD Keywords list MUST appear somewhere in the resume, exactly
as written in the JD (case matters: "Next.js" not "NextJS", "Node.js" not "NodeJS").
Weave them naturally — do not just dump them in the skills section.

RULE 3 — ROLE-TYPE EMPHASIS:
Follow the ROLE TYPE EMPHASIS instructions above strictly.
The experience sections must be REORDERED and REWEIGHTED — not just relabelled.

RULE 4 — MIRROR COMPANY LANGUAGE:
The Professional Summary MUST use words and phrases from the company's culture page
or JD — their mission language, their values vocabulary, their product terminology.

RULE 5 — EVERY BULLET = RESULT:
No bullet can start with "Responsible for" or "Helped with" or "Worked on".
Every bullet: Action verb + What was done + Measurable outcome or scale.
Use real metrics from the bank (e.g. "500k+ contacts", "40% page load reduction",
"6 brands", "157 TypeScript errors", "634MB vector DB").

=== OUTPUT FORMAT — PLAIN TEXT SECTIONS (no LaTeX, no markdown fences) ===

Return EXACTLY two blocks separated by this line:
JOBCRAFT_SPLIT_DELIMITER

Block 1 (BEFORE delimiter): Resume content using EXACTLY these section tags:

=== SUMMARY ===
[4-5 sentences: who you are + role title verbatim | company language mirror |
strongest metric achievement | secondary strength | unique differentiator]

=== SKILLS ===
[Each line: "Category: skill1, skill2, skill3" — ordered by JD relevance.
Groups: Languages | Frameworks | Databases | MarTech & Email | eCommerce | Tools & Platforms
Only include skills relevant to this role. Every JD keyword that maps to a skill must appear here.]

=== EXPERIENCE ===
[Company Name — Title | Approx Dates]
- Bullet (Action verb + what + measurable result)
- Bullet
[Next company...]

=== PROJECTS ===
[Project Name | Tech Stack (JD-relevant first)]
- Impact bullet
[Max 3 projects, chosen by relevance to this role]

=== EDUCATION ===
[Whatever is in the candidate profile JSON — do not fabricate degrees]

Block 2 (AFTER delimiter): Markdown alignment audit:

# Alignment Audit — {company} ({title})
## Role Type Detected
{role_type}

## ATS Keyword Coverage
| JD Keyword | In Resume | Source |
|---|---|---|
[One row per keyword. ✅ if present, ❌ if genuinely absent from experience bank]

**Estimated ATS Match Score: XX/100**

## What Was Emphasized and Why
[2-3 sentences explaining which experiences led the resume and why, per role type logic]

## Genuine Skill Gaps
[Honest list of JD requirements not in experience bank. If none: "No material gaps."]

## Company Culture Mirror
[Quote 2-3 phrases from the company culture page and show where they appear in the resume.
If no culture page was found, note that and explain how JD language was used instead.]

## Interview Preparation Playbook

### Recruiter / HR Screen
- Questions to expect (specific to this company and role):
- Your anchor stories (which experience to cite):
- Salary anchor: [realistic range for this role/company tier]

### Technical Screen
- Topics to review (specific to JD, not generic):
- Practice problems (realistic for this company type):
- Format expected:

### Technical Deep Dive / System Design
- Realistic design scenario for this company's domain:
- Architecture talking points from your experience:
- Technologies to demonstrate from JD:

### Culture / Behavioral Round
- Company values to address (from culture page or JD):
- STAR stories to prepare (map your real experiences):
- Smart questions to ask:

### Your Strongest Differentiator for This Role
[One paragraph: what makes Yogita genuinely different from other candidates applying here]

## Cover Letter Hook
[2-3 sentence opening paragraph for a cover letter or cold email to the hiring manager.
Uses company's own language. References a specific detail from the JD or culture page.
Ends with a clear hook that makes them want to read more.]

## Sources
- Experience Bank: resume_builder.py EXPERIENCE_BANK
- Company Culture: {job.get('url', 'N/A')} (and About/Careers page if scraped)
- JD Source: {job.get('url', 'N/A')}
- Generated by: JobCraft AI Resume Engine
"""

    print(f"🚀 Generating tailored Word resume for: {company} — {title} [{role_type}]")

    max_retries = 3
    retry_delay = 65

    def parse_sections(text: str) -> dict:
        """Extracts tagged sections from Gemini plain-text output."""
        tags = ['SUMMARY', 'SKILLS', 'EXPERIENCE', 'PROJECTS', 'EDUCATION']
        result = {}
        for i, tag in enumerate(tags):
            start_marker = f"=== {tag} ==="
            end_markers  = [f"=== {t} ===" for t in tags[i + 1:]]
            start = text.find(start_marker)
            if start == -1:
                continue
            start += len(start_marker)
            end = len(text)
            for em in end_markers:
                pos = text.find(em, start)
                if pos != -1 and pos < end:
                    end = pos
            result[tag.lower()] = text[start:end].strip()
        return result

    for attempt in range(1, max_retries + 1):
        try:
            response = llm.invoke([prompt])
            content  = response.content.strip()

            # Strip outer markdown fences if present
            content = re.sub(r'^```[a-z]*\n', '', content)
            content = re.sub(r'\n```$', '', content)

            # Split on delimiter
            parts = re.split(r'\s*JOBCRAFT_SPLIT_DELIMITER\s*', content, maxsplit=1)

            if len(parts) == 2:
                resume_text     = parts[0].strip()
                alignment_audit = parts[1].strip()
            else:
                # Fallback: assume everything before first "# Alignment Audit" is resume
                split_idx = content.find('# Alignment Audit')
                if split_idx != -1:
                    resume_text     = content[:split_idx].strip()
                    alignment_audit = content[split_idx:].strip()
                else:
                    resume_text     = content
                    alignment_audit = f"# Alignment Audit — {company} ({title})\n\n⚠️ Split delimiter not found."

            # Clean audit fences
            alignment_audit = re.sub(r'^```markdown\s*\n?', '', alignment_audit)
            alignment_audit = re.sub(r'\n?```\s*$', '', alignment_audit)

            sections = parse_sections(resume_text)
            if not sections.get('summary') and not sections.get('experience'):
                raise ValueError("No recognisable sections found in Gemini output — retrying.")

            # Build Word document
            build_docx(sections, company, title, docx_path)

            with open(md_path, 'w', encoding='utf-8') as f:
                f.write(alignment_audit)
            print(f"📋 Alignment audit written: {md_path}")

            return docx_path, md_path

        except ValueError as ve:
            print(f"⚠️  Attempt {attempt}/{max_retries} — {ve}")
            if attempt < max_retries:
                print(f"   Retrying in {retry_delay}s...")
                time.sleep(retry_delay)
            continue

        except Exception as e:
            err = str(e)
            if "429" in err or "RESOURCE_EXHAUSTED" in err or "quota" in err.lower():
                delay_match = re.search(r'retry[_\s]?delay[:\s]+[\'"]?(\d+)', err, re.IGNORECASE)
                wait = int(delay_match.group(1)) + 5 if delay_match else retry_delay
                print(f"⚠️  Attempt {attempt}/{max_retries} — Gemini quota (429). Waiting {wait}s...")
                if attempt < max_retries:
                    time.sleep(wait)
                continue
            else:
                print(f"❌ Attempt {attempt}/{max_retries} — Unexpected error: {e}")
                if attempt < max_retries:
                    time.sleep(10)
                continue

    # All retries exhausted — write minimal fallback docx
    print(f"❌ All {max_retries} attempts failed. Writing diagnostic fallback.")

    fallback_sections = {
        'summary': (
            f"Resume generation failed after {max_retries} retries. "
            f"Likely cause: Gemini API free-tier quota exhausted. "
            f"Re-run when quota resets (~24h). Role: {title} at {company}."
        ),
        'skills':     f"Role type detected: {role_type}\nJD keywords: {keywords_str}",
        'experience': "See job URL below for role details.",
        'projects':   "",
        'education':  "Refer to candidate_profile.json"
    }
    build_docx(fallback_sections, company, title, docx_path)

    fallback_md = (
        f"# Alignment Audit — {company} ({title})\n\n"
        f"> ⚠️ **Generation failed after {max_retries} retries.**\n\n"
        f"**Likely cause:** Gemini API free-tier quota exhausted.\n\n"
        f"**What to do:**\n"
        f"1. Wait 24h for quota reset, then re-run.\n"
        f"2. Add `GOOGLE_API_KEY2` to .env for key rotation.\n"
        f"3. Check: https://ai.dev/rate-limit\n\n"
        f"**Role type detected:** {role_type}\n"
        f"**JD keywords found:** {keywords_str}\n"
        f"**Job URL:** {job.get('url', 'N/A')}\n"
    )
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(fallback_md)

    return docx_path, md_path